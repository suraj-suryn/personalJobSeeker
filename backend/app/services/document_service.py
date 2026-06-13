"""
PDF and DOCX document generation for resumes and cover letters.
"""

import logging
import uuid
from pathlib import Path

from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


def _get_output_path(subdir: str, filename: str) -> Path:
    path = Path(settings.generated_dir) / subdir
    path.mkdir(parents=True, exist_ok=True)
    return path / filename


# ─────────────────────────────────────────────
#  PDF Generation (reportlab)
# ─────────────────────────────────────────────

def generate_resume_pdf(content: dict, output_id: str) -> str:
    """Generate a PDF resume from structured content dict. Returns file path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    file_path = _get_output_path("resumes", f"{output_id}.pdf")

    doc = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", fontSize=18, fontName="Helvetica-Bold", spaceAfter=4)
    contact_style = ParagraphStyle("Contact", fontSize=10, textColor=colors.grey, spaceAfter=12)
    section_style = ParagraphStyle("Section", fontSize=12, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=4)
    body_style = ParagraphStyle("Body", fontSize=10, spaceAfter=4, leading=14)
    skill_style = ParagraphStyle("Skills", fontSize=10, spaceAfter=2)

    story = []

    # Header
    name = content.get("name") or "Name"
    story.append(Paragraph(name, name_style))

    contact_parts = []
    if content.get("email"):
        contact_parts.append(content["email"])
    if content.get("phone"):
        contact_parts.append(content["phone"])
    if content.get("location"):
        contact_parts.append(content["location"])
    if content.get("linkedin_url"):
        contact_parts.append(content["linkedin_url"])
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))

    # Summary
    if content.get("summary"):
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2563eb")))
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        story.append(Paragraph(content["summary"], body_style))

    # Skills
    if content.get("skills"):
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2563eb")))
        story.append(Paragraph("SKILLS", section_style))
        skills_text = " • ".join(content["skills"])
        story.append(Paragraph(skills_text, skill_style))

    # Experience
    if content.get("experience"):
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2563eb")))
        story.append(Paragraph("EXPERIENCE", section_style))
        for exp in content["experience"]:
            title_company = f"<b>{exp.get('title', '')}</b> — {exp.get('company', '')}"
            story.append(Paragraph(title_company, body_style))
            date_loc = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')} | {exp.get('location', '')}"
            story.append(Paragraph(date_loc, ParagraphStyle("DateLoc", fontSize=9, textColor=colors.grey)))
            if exp.get("description"):
                story.append(Paragraph(exp["description"], body_style))
            for ach in exp.get("achievements", []):
                story.append(Paragraph(f"• {ach}", body_style))
            story.append(Spacer(1, 6))

    # Education
    if content.get("education"):
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2563eb")))
        story.append(Paragraph("EDUCATION", section_style))
        for edu in content["education"]:
            edu_text = f"<b>{edu.get('degree', '')}</b>, {edu.get('institution', '')}"
            story.append(Paragraph(edu_text, body_style))
            year_text = f"{edu.get('start_year', '')} – {edu.get('end_year', '')}"
            story.append(Paragraph(year_text, ParagraphStyle("Year", fontSize=9, textColor=colors.grey)))

    # Certifications
    if content.get("certifications"):
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2563eb")))
        story.append(Paragraph("CERTIFICATIONS", section_style))
        for cert in content["certifications"]:
            cert_text = f"• {cert.get('name', '')} — {cert.get('issuer', '')} ({cert.get('year', '')})"
            story.append(Paragraph(cert_text, body_style))

    doc.build(story)
    return str(file_path)


def generate_cover_letter_pdf(content: str, output_id: str) -> str:
    """Generate a cover letter PDF. Returns file path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    file_path = _get_output_path("cover_letters", f"{output_id}.pdf")
    doc = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        leftMargin=3 * cm,
        rightMargin=3 * cm,
        topMargin=3 * cm,
        bottomMargin=3 * cm,
    )

    body_style = ParagraphStyle("Body", fontSize=11, leading=16, spaceAfter=12)
    story = []
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            story.append(Paragraph(paragraph.strip().replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 6))

    doc.build(story)
    return str(file_path)


# ─────────────────────────────────────────────
#  DOCX Generation (python-docx)
# ─────────────────────────────────────────────

def generate_resume_docx(content: dict, output_id: str) -> str:
    """Generate a DOCX resume. Returns file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    file_path = _get_output_path("resumes", f"{output_id}.docx")
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Name header
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(content.get("name") or "Name")
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(37, 99, 235)  # blue-600

    # Contact
    contact_parts = [
        content.get("email", ""),
        content.get("phone", ""),
        content.get("location", ""),
        content.get("linkedin_url", ""),
    ]
    contact_para = doc.add_paragraph(" | ".join(p for p in contact_parts if p))
    contact_para.runs[0].font.size = Pt(9)
    contact_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    def add_section(title: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        doc.add_paragraph("─" * 60)

    # Summary
    if content.get("summary"):
        add_section("PROFESSIONAL SUMMARY")
        doc.add_paragraph(content["summary"])

    # Skills
    if content.get("skills"):
        add_section("SKILLS")
        doc.add_paragraph(" • ".join(content["skills"]))

    # Experience
    if content.get("experience"):
        add_section("EXPERIENCE")
        for exp in content["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            run.font.bold = True
            date_p = doc.add_paragraph(
                f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')} | {exp.get('location', '')}"
            )
            date_p.runs[0].font.size = Pt(9)
            if exp.get("description"):
                doc.add_paragraph(exp["description"])
            for ach in exp.get("achievements", []):
                doc.add_paragraph(f"• {ach}", style="List Bullet")

    # Education
    if content.get("education"):
        add_section("EDUCATION")
        for edu in content["education"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{edu.get('degree', '')}, {edu.get('institution', '')}")
            r.font.bold = True
            doc.add_paragraph(f"{edu.get('start_year', '')} – {edu.get('end_year', '')}")

    doc.save(str(file_path))
    return str(file_path)


def generate_cover_letter_docx(content: str, output_id: str) -> str:
    """Generate a cover letter DOCX. Returns file path."""
    from docx import Document
    from docx.shared import Pt

    file_path = _get_output_path("cover_letters", f"{output_id}.docx")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            p.paragraph_format.space_after = Pt(12)

    doc.save(str(file_path))
    return str(file_path)
