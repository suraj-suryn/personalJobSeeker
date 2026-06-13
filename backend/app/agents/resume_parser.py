"""
Resume Parser Agent
Extracts structured data from PDF/DOCX files using LLM.
"""

import json
import logging
from pathlib import Path

from app.core.llm import get_llm
from app.schemas.resume import ParsedResumeData

logger = logging.getLogger(__name__)


class ResumeParser:
    def __init__(self) -> None:
        self.llm = get_llm()

    async def parse(
        self, file_path: str, ext: str, provider: str | None = None
    ) -> tuple[str, ParsedResumeData | None]:
        """
        Extract text from file, then use LLM to structure it.
        Returns (raw_text, ParsedResumeData).
        """
        raw_text = self._extract_text(file_path, ext)
        if not raw_text.strip():
            return "", None

        parsed = await self._llm_extract(raw_text, provider=provider)
        return raw_text, parsed

    def _extract_text(self, file_path: str, ext: str) -> str:
        """Extract plain text from PDF or DOCX."""
        ext = ext.lower().lstrip(".")
        try:
            if ext == "pdf":
                return self._extract_pdf(file_path)
            elif ext == "docx":
                return self._extract_docx(file_path)
        except Exception as exc:
            logger.warning("Text extraction failed for %s: %s", file_path, exc)
        return ""

    def _extract_pdf(self, file_path: str) -> str:
        # Try pdfminer first (better layout), fallback to PyPDF2
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(file_path)
            if text and text.strip():
                return text
        except Exception:
            pass

        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
        except Exception as exc:
            logger.error("PDF extraction failed: %s", exc)
            return ""

    def _extract_docx(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    async def _llm_extract(
        self, raw_text: str, provider: str | None = None
    ) -> ParsedResumeData | None:
        """Use LLM to extract structured data from resume text."""

        system_prompt = """You are a resume parsing assistant. Extract structured data from the resume text provided.
Return ONLY valid JSON matching this exact schema (no markdown, no explanation):
{
  "name": "string or null",
  "email": "string or null",
  "phone": "string or null",
  "location": "string or null",
  "linkedin_url": "string or null",
  "github_url": "string or null",
  "summary": "string or null",
  "skills": ["skill1", "skill2"],
  "education": [{"institution": "", "degree": "", "field_of_study": "", "start_year": "", "end_year": "", "gpa": ""}],
  "experience": [{"company": "", "title": "", "location": "", "start_date": "", "end_date": "", "is_current": false, "description": "", "achievements": []}],
  "certifications": [{"name": "", "issuer": "", "year": "", "url": ""}],
  "projects": [{"name": "", "description": "", "technologies": [], "url": ""}],
  "languages": []
}
Extract all skills mentioned. Be thorough with achievements in experience."""

        user_prompt = f"Parse this resume:\n\n{raw_text[:6000]}"

        try:
            response = await self.llm.chat(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                provider=provider,
                temperature=0.1,
                max_tokens=4096,
            )

            # Clean and parse JSON
            cleaned = response.strip()
            if cleaned.startswith("```"):
                lines = cleaned.split("\n")
                cleaned = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])

            data = json.loads(cleaned)
            return ParsedResumeData(**data)

        except json.JSONDecodeError as exc:
            logger.warning("LLM returned invalid JSON for resume: %s", exc)
            # Best-effort: return minimal data
            return ParsedResumeData()
        except Exception as exc:
            logger.exception("LLM resume extraction failed: %s", exc)
            return None
